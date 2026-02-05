from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_article_front(doc, title, authors):
    # Espacio después del header
    doc.add_paragraph("")

    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("")

    # Autores
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(authors)
    run.font.size = Pt(10)

    doc.add_paragraph("")
