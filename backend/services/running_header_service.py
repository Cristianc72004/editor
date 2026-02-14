# backend/services/running_header_service.py

from docx.shared import Inches
import os, time

from .header_utils import (
    _ensure_generated_dir,
    _add_floating_picture,
    _build_bar_with_notch_png,
    _inline_to_anchor,
    _pick_font,
    _text_size,
    GREEN_HEX,
    WHITE_RGB
)

from PIL import Image, ImageDraw


def _build_textbox_png(path_png, width_in, height_in, text, font_pt=10, dpi=220):
    W = int(round(width_in * dpi))
    H = int(round(height_in * dpi))

    img = Image.new("RGBA", (W, H), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)

    font = _pick_font(int(font_pt * dpi / 72), bold=False)

    tw, th = _text_size(draw, text, font)

    x = W - tw - 5   # alineado a la derecha dentro del cuadro
    y = (H - th) // 2

    draw.text((x, y), text, fill=(0, 0, 0, 255), font=font)

    img.save(path_png, format="PNG")


def add_running_header(
    header,
    section,
    review,
    author,

    review_x,
    review_y,
    review_w,
    review_h,

    author_x,
    author_y,
    author_w,
    author_h,

    line_x,
    line_y,
    line_w,
    line_h
):

    # Limpiar header
    while header.paragraphs:
        p = header.paragraphs[0]._p
        header._element.remove(p)

    out_dir = _ensure_generated_dir()
    ts = str(int(time.time() * 1000))

    # =============================
    # 1️⃣ CUADRO REVIEW
    # =============================

    review_png = os.path.join(out_dir, f"running_review_box_{ts}.png")

    _build_bar_with_notch_png(
        review_png,
        review_w,
        review_h,
        GREEN_HEX,
        f"  {review}  ",
        10,
        WHITE_RGB
    )

    p_box = header.add_paragraph()

    _add_floating_picture(
        p_box,
        review_png,
        review_w,
        review_x,
        review_y,
        h_ref="page",
        v_ref="page",
        z_index=2000,
        behind_doc=False
    )

    # =============================
    # 2️⃣ AUTOR COMO CUADRO FLOTANTE
    # =============================

    author_png = os.path.join(out_dir, f"running_author_{ts}.png")

    _build_textbox_png(
        author_png,
        author_w,
        author_h,
        author,
        font_pt=10
    )

    p_author = header.add_paragraph()

    _add_floating_picture(
        p_author,
        author_png,
        author_w,
        author_x,
        author_y,
        h_ref="page",
        v_ref="page",
        z_index=3000,
        behind_doc=False
    )

    # =============================
    # 3️⃣ LÍNEA INFERIOR
    # =============================

    line_png = os.path.join(out_dir, f"running_line_{ts}.png")

    _build_bar_with_notch_png(
        line_png,
        line_w,
        line_h,
        GREEN_HEX,
        "",
        1,
        WHITE_RGB
    )

    p_line = header.add_paragraph()

    _add_floating_picture(
        p_line,
        line_png,
        line_w,
        line_x,
        line_y,
        h_ref="page",
        v_ref="page",
        z_index=1000,
        behind_doc=True
    )
