# backend/services/document_service.py

import os, time
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_SECTION

from .first_page_header_service import add_first_page_header
from .running_header_service import add_running_header
from .footer_service import add_footer_section
from .article_service import add_article_front

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
DOC_ORIGINAL = os.path.join(BASE_DIR, "documents", "original")
DOC_PROCESSED = os.path.join(BASE_DIR, "documents", "processed")
LOGOS_DIR = os.path.join(BASE_DIR, "assets", "logos")

REVIEW_TYPE = "Review"


def list_documents():
    return [
        f for f in os.listdir(DOC_ORIGINAL)
        if f.endswith(".docx")
    ]


def process_document(
    filename,
    logo_left,
    logo_right,
    running_author,
    title,
    authors,
    footer_text,
    journal_name,
    issn,

    logo_left_x, logo_left_y, logo_left_w, logo_left_h,
    logo_right_x, logo_right_y, logo_right_w, logo_right_h,
    title_x, title_y, title_w, title_h,
    bar_x, bar_y, bar_w, bar_h,

    review_x=0.4,
    review_y=0.3,
    review_w=1.2,
    review_h=0.35,

    author_x=6.8,
    author_y=0.35,
    author_w=1.0,
    author_h=0.3,

    line_x=0.4,
    line_y=0.75,
    line_w=7.4,
    line_h=0.05
):

    input_path = os.path.join(DOC_ORIGINAL, filename)

    os.makedirs(DOC_PROCESSED, exist_ok=True)
    ts = int(time.time())

    output_path = os.path.join(
        DOC_PROCESSED,
        f"{os.path.splitext(filename)[0]}_editado_{ts}.docx"
    )

    original = Document(input_path)
    doc = Document()

    # ==================================================
    # SECCIÓN 1 → SOLO PRIMERA PÁGINA
    # ==================================================
    section1 = doc.sections[0]

    section1.page_width = Inches(8.27)
    section1.page_height = Inches(11.69)

    section1.top_margin = Inches(2.2)
    section1.header_distance = Inches(0.2)

    section1.bottom_margin = Inches(1)
    section1.footer_distance = Inches(0.3)

    # HEADER GRANDE
    add_first_page_header(
        header=section1.header,
        section=section1,
        logo_left=os.path.join(LOGOS_DIR, logo_left),
        logo_right=os.path.join(LOGOS_DIR, logo_right)
        if logo_right else None,
        review=REVIEW_TYPE,
        journal_name=journal_name,
        issn=issn,

        logo_left_x=logo_left_x,
        logo_left_y=logo_left_y,
        logo_left_w=logo_left_w,
        logo_left_h=logo_left_h,

        logo_right_x=logo_right_x,
        logo_right_y=logo_right_y,
        logo_right_w=logo_right_w,
        logo_right_h=logo_right_h,

        title_x=title_x,
        title_y=title_y,
        title_w=title_w,
        title_h=title_h,

        bar_x=bar_x,
        bar_y=bar_y,
        bar_w=bar_w,
        bar_h=bar_h
    )

    # FOOTER sección 1
    add_footer_section(section1, footer_text)

    # ==================================================
    # CREAR SECCIÓN 2 → RESTO DEL DOCUMENTO
    # ==================================================
    section2 = doc.add_section(WD_SECTION.NEW_PAGE)

    section2.page_width = Inches(8.27)
    section2.page_height = Inches(11.69)

    section2.top_margin = Inches(1.1)
    section2.header_distance = Inches(0.2)

    section2.bottom_margin = Inches(1)
    section2.footer_distance = Inches(0.3)

    section2.header.is_linked_to_previous = False

    # HEADER PEQUEÑO
    add_running_header(
        header=section2.header,
        section=section2,
        review=REVIEW_TYPE,

        review_x=review_x,
        review_y=review_y,
        review_w=review_w,
        review_h=review_h,

        author=running_author,
        author_x=author_x,
        author_y=author_y,
        author_w=author_w,
        author_h=author_h,

        line_x=line_x,
        line_y=line_y,
        line_w=line_w,
        line_h=line_h
    )

    # FOOTER sección 2
    add_footer_section(section2, footer_text)

    # ==================================================
    # AHORA SÍ → CONTENIDO ORIGINAL VA A SECCIÓN 2
    # ==================================================
    for p in original.paragraphs:
        doc.add_paragraph(p.text)

    # FRONT MATTER
    add_article_front(
        doc=doc,
        title=title,
        authors=authors
    )

    doc.save(output_path)
    return output_path
