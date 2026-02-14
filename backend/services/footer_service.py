# backend/services/footer_service.py

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_footer_section(section, text):
    """
    Footer por sección.
    No usa doc.sections[-1].
    Evita duplicados.
    """

    footer = section.footer

    # limpiar footer
    while footer.paragraphs:
        p = footer.paragraphs[0]._p
        footer._element.remove(p)

    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p.add_run(text)
    p.add_run("\t")

    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
