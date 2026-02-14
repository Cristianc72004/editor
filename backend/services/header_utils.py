# backend/services/header_utils.py

from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import os, time

GREEN_HEX = "1F5C4D"
GREY_ISSN_HEX = "557075"
WHITE_RGBA = (255, 255, 255, 255)
WHITE_RGB = (255, 255, 255)
TRANSPARENT = (255, 255, 255, 0)
BLACK_RGB = (0, 0, 0)

_EMUS_PER_INCH = 914400
def _emus(inches: float) -> int:
    return int(round(inches * _EMUS_PER_INCH))

def _ensure_dir(path: str): os.makedirs(path, exist_ok=True)

def _ensure_generated_dir():
    base = os.path.dirname(os.path.dirname(__file__))
    out = os.path.join(base, "assets", "generated")
    _ensure_dir(out)
    return out

def _pick_font(px: int, bold: bool = False):
    try:
        candidates = []
        if bold:
            candidates += [
                "C:/Windows/Fonts/timesbd.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSerif-Bold.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
            ]
        else:
            candidates += [
                "C:/Windows/Fonts/times.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
            ]
        candidates += [
            "C:/Windows/Fonts/arial.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ]
        for fp in candidates:
            if os.path.exists(fp):
                return ImageFont.truetype(fp, px)
        return ImageFont.load_default()
    except:
        return ImageFont.load_default()

def _text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont):
    l, t, r, b = draw.textbbox((0, 0), text, font=font)
    return r - l, b - t

def _hex_to_rgb(h: str):
    h = h.strip().lstrip("#")
    return (int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def _inline_to_anchor(inline_el, pos_x_in: float, pos_y_in: float,
                      h_ref="page", v_ref="page",
                      z_index=2000, behind_doc: bool=False):
    extent = inline_el.find(qn("wp:extent"))
    cx = extent.get("cx") if extent is not None else "0"
    cy = extent.get("cy") if extent is not None else "0"

    docPr = inline_el.find(qn("wp:docPr"))
    cNv = inline_el.find(qn("wp:cNvGraphicFramePr"))
    graphic = inline_el.find(qn("a:graphic"))

    anchor = OxmlElement("wp:anchor")
    for k, v in {
        "relativeHeight": str(z_index),
        "distT": "0", "distB": "0", "distL": "0", "distR": "0",
        "simplePos": "0",
        "behindDoc": "1" if behind_doc else "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "1",
    }.items(): anchor.set(k, v)

    simplePos = OxmlElement("wp:simplePos")
    simplePos.set("x", "0"); simplePos.set("y", "0")
    anchor.append(simplePos)

    positionH = OxmlElement("wp:positionH")
    positionH.set("relativeFrom", h_ref)
    posH = OxmlElement("wp:posOffset"); posH.text = str(_emus(pos_x_in))
    positionH.append(posH)
    anchor.append(positionH)

    positionV = OxmlElement("wp:positionV")
    positionV.set("relativeFrom", v_ref)
    posV = OxmlElement("wp:posOffset"); posV.text = str(_emus(pos_y_in))
    positionV.append(posV)
    anchor.append(positionV)

    new_extent = OxmlElement("wp:extent")
    new_extent.set("cx", cx)
    new_extent.set("cy", cy)
    anchor.append(new_extent)

    effectExtent = OxmlElement("wp:effectExtent")
    for s in ["l", "t", "r", "b"]: effectExtent.set(s, "0")
    anchor.append(effectExtent)

    wrap = OxmlElement("wp:wrapNone") if behind_doc else OxmlElement("wp:wrapSquare")
    if not behind_doc: wrap.set("wrapText", "bothSides")
    anchor.append(wrap)

    if docPr is not None: anchor.append(docPr)
    if cNv   is not None: anchor.append(cNv)
    if graphic is not None: anchor.append(graphic)

    parent = inline_el.getparent(); idx = parent.index(inline_el)
    parent.remove(inline_el); parent.insert(idx, anchor)

def _add_floating_picture(paragraph, image_path, width_in, x_in, y_in,
                          height_in=None,
                          h_ref="page", v_ref="page",
                          z_index=2000, behind_doc: bool=False):
    run = paragraph.add_run()
    if height_in is None or height_in <= 0:
        run.add_picture(image_path, width=Inches(width_in))
    else:
        run.add_picture(image_path, width=Inches(width_in), height=Inches(height_in))
    inl = run._r.xpath(".//wp:inline")
    if not inl: return
    _inline_to_anchor(inl[-1], x_in, y_in, h_ref=h_ref, v_ref=v_ref, z_index=z_index, behind_doc=behind_doc)

def _build_title_png(path_png, width_in, height_in, title, issn, title_pt, issn_pt, gap_px=18, dpi=220):
    W = max(40, int(round(width_in  * dpi)))
    H = max(10, int(round(height_in * dpi)))
    img = Image.new("RGBA", (W, H), TRANSPARENT)
    draw = ImageDraw.Draw(img)

    ft = _pick_font(int(title_pt * dpi / 72), bold=True)
    fi = _pick_font(int(issn_pt  * dpi / 72), bold=False)

    issn_val = issn.replace("ISSN", "").replace(":", "").strip()
    label = "ISSN: "

    tw, _ = _text_size(draw, title, ft)
    lw, _ = _text_size(draw, label, fi)
    iw, ih = _text_size(draw, issn_val, fi)

    total_w = tw + gap_px + lw + iw
    if total_w > W - 8:
        gap_px = max(10, gap_px - 6)
        total_w = tw + gap_px + lw + iw

    ft_ascent, ft_descent = ft.getmetrics()
    fi_ascent, fi_descent = fi.getmetrics()
    max_ascent = max(ft_ascent, fi_ascent)
    max_descent = max(ft_descent, fi_descent)
    line_h = max_ascent + max_descent

    y0 = max(0, (H - line_h) // 2)
    x = max(0, (W - total_w) // 2)

    green = _hex_to_rgb(GREEN_HEX)
    grey = _hex_to_rgb(GREY_ISSN_HEX)

    y_title = y0 + (max_ascent - ft_ascent)
    draw.text((x, y_title), title, fill=green + (255,), font=ft)
    x += tw + gap_px

    y_issn = y0 + (max_ascent - fi_ascent)
    draw.text((x, y_issn), label,    fill=grey + (255,), font=fi)
    x += lw
    draw.text((x, y_issn), issn_val, fill=grey + (255,), font=fi)

    img.save(path_png, format="PNG")

def _build_bar_with_notch_png(path_png, width_in, height_in, fill_hex, text, text_pt, text_color=WHITE_RGB, notch_w_in=0.18, notch_h_in=0.18, dpi=220):
    W = max(20, int(round(width_in  * dpi)))
    H = max(10, int(round(height_in * dpi)))
    img = Image.new("RGBA", (W, H), TRANSPARENT)
    draw = ImageDraw.Draw(img)

    fill = _hex_to_rgb(fill_hex)
    draw.rectangle([0, 0, W, H], fill=fill + (255,))

    nw = int(round(notch_w_in * dpi))
    nh = int(round(notch_h_in * dpi))
    tri = [(W - nw, 0), (W, 0), (W, nh)]
    draw.polygon(tri, fill=WHITE_RGBA)

    f = _pick_font(int(text_pt * dpi / 72))
    _, th = _text_size(draw, text, f)
    y = (H - th) // 2
    padL = max(10, int(H * 0.8))
    draw.text((padL, y), text, fill=text_color + (255,), font=f)

    img.save(path_png, format="PNG")

def _probe_image_size_in(image_path: str):
    try:
        from PIL import Image as PILImage
        with PILImage.open(image_path) as im:
            w_px, h_px = im.size
            dpi = None
            if "dpi" in im.info and isinstance(im.info["dpi"], (tuple, list)) and len(im.info["dpi"]) >= 1:
                dpi = im.info["dpi"][0]
            if not dpi or dpi <= 0:
                dpi = 96.0
            return w_px / dpi, h_px / dpi
    except:
        return (1.0, 1.0)

def _size_from_w_or_h(image_path: str, w_in: float, h_in: float):
    nat_w, nat_h = _probe_image_size_in(image_path)
    if nat_w <= 0 or nat_h <= 0:
        return (w_in, w_in * 0.8) if h_in <= 0 else (h_in, h_in)
    ratio = nat_h / nat_w
    if h_in <= 0:
        return w_in, w_in * ratio
    else:
        if ratio == 0:
            return w_in, h_in
        new_w = h_in / ratio
        return new_w, h_in
