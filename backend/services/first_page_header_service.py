# backend/services/first_page_header_service.py

from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, time

from .header_utils import (
    _ensure_generated_dir,
    _size_from_w_or_h,
    _build_bar_with_notch_png,
    _add_floating_picture,
    _inline_to_anchor,
    _build_title_png,
    GREEN_HEX,
    WHITE_RGB
)


def add_first_page_header(
    header, section,
    logo_left, logo_right,
    review, journal_name, issn,
    logo_left_x, logo_left_y, logo_left_w, logo_left_h,
    logo_right_x, logo_right_y, logo_right_w, logo_right_h,
    title_x, title_y, title_w, title_h,
    bar_x, bar_y, bar_w, bar_h
):
    # Limpia header
    while header.paragraphs:
        p = header.paragraphs[0]._p
        header._element.remove(p)

    out_dir = _ensure_generated_dir()
    ts = str(int(time.time() * 1000))

    page_w_in = section.page_width.inches

    # Tamaños finales para logos
    left_w, left_h = _size_from_w_or_h(logo_left, logo_left_w, logo_left_h)
    right_w, right_h = (0.0, 0.0)
    if logo_right:
        right_w, right_h = _size_from_w_or_h(logo_right, logo_right_w, logo_right_h)

    # (1) FRANJA (detrás)
    bar_png = os.path.join(out_dir, f"bar_review_notch_{ts}.png")
    _build_bar_with_notch_png(
        bar_png,
        bar_w,
        bar_h,
        GREEN_HEX,
        f"  {review}  ",
        11,
        WHITE_RGB
    )

    p_bar = header.add_paragraph()
    _add_floating_picture(
        p_bar,
        bar_png,
        bar_w,
        bar_x,
        bar_y,
        h_ref="page",
        v_ref="page",
        z_index=500,
        behind_doc=True
    )

    # (2) LOGO IZQ
    p1 = header.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rL = p1.add_run()

    if logo_left_h <= 0:
        rL.add_picture(logo_left, width=Inches(left_w))
    else:
        rL.add_picture(
            logo_left,
            width=Inches(left_w),
            height=Inches(left_h)
        )

    inlL = rL._r.xpath(".//wp:inline")
    if inlL:
        _inline_to_anchor(
            inlL[-1],
            logo_left_x,
            logo_left_y,
            "page",
            "page",
            2000,
            False
        )

    # (3) LOGO DER
    if logo_right:
        p2 = header.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rR = p2.add_run()

        if logo_right_h <= 0:
            rR.add_picture(logo_right, width=Inches(right_w))
        else:
            rR.add_picture(
                logo_right,
                width=Inches(right_w),
                height=Inches(right_h)
            )

        inlR = rR._r.xpath(".//wp:inline")
        if inlR:
            RIGHT_PADDING = 0.10
            x_right = (
                max(0.0, page_w_in - right_w - RIGHT_PADDING)
                if logo_right_x < 0
                else logo_right_x
            )

            _inline_to_anchor(
                inlR[-1],
                x_right,
                logo_right_y,
                "page",
                "page",
                2000,
                False
            )

    # (4) TÍTULO + ISSN
    title_png = os.path.join(out_dir, f"title_box_{ts}.png")

    _build_title_png(
        title_png,
        title_w,
        title_h,
        journal_name,
        issn,
        24,
        12,
        18,
        220
    )

    p3 = header.add_paragraph()

    _add_floating_picture(
        p3,
        title_png,
        title_w,
        title_x,
        title_y,
        h_ref="page",
        v_ref="page",
        z_index=3000,
        behind_doc=False
    )
